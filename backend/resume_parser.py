import pdfplumber
from docx import Document
import os
from utils.candidate_classifier import detect_candidate_type


def extract_text(file_path):

    ext = os.path.splitext(file_path)[1].lower()

    text = ""


    if ext == ".pdf":

        with pdfplumber.open(file_path) as pdf:

            for page in pdf.pages:

                page_text = page.extract_text()

                if page_text:
                    text += page_text + "\n"



    elif ext == ".docx":

        doc = Document(file_path)


        # Paragraphs
        for para in doc.paragraphs:
            text += para.text + "\n"


        # Tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\n"



    else:
        raise ValueError("Unsupported file format")


   
    return text